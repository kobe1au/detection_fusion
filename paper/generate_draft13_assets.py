from __future__ import annotations

import html
import re
import subprocess
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\code\detection_fusion_copy\paper\定稿13_统一图表")
SVG_DIR = ROOT / "svg"
PNG_DIR = ROOT / "png"
EDGE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")

W = 1600
FONT_CN = "SimSun, 'Microsoft YaHei', sans-serif"
FONT_EN = "'Times New Roman', serif"
STROKE = "#333333"
GRAY = "#F4F4F4"
API = "#EAF4E5"
GRAPH = "#E8F0FA"
MAN = "#FFF1DF"
ACCENT = "#EDEDED"
REJECT = "#F8ECEC"


def esc(value: str) -> str:
    return html.escape(value, quote=True)


def svg_start(height: int) -> list[str]:
    return [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{height}" viewBox="0 0 {W} {height}">',
        '<defs>',
        '<marker id="arrow" markerWidth="10" markerHeight="10" refX="8.5" refY="3.5" orient="auto" markerUnits="strokeWidth">',
        f'<path d="M0,0 L9,3.5 L0,7 Z" fill="{STROKE}"/>',
        '</marker>',
        '</defs>',
        '<rect width="100%" height="100%" fill="#FFFFFF"/>',
    ]


def svg_end(parts: list[str]) -> str:
    parts.append('</svg>')
    return '\n'.join(parts)


def box(parts: list[str], x: float, y: float, w: float, h: float, lines: list[str], *, fill: str = "#FFFFFF", stroke: str = STROKE, radius: int = 8, font_size: int = 28, bold_first: bool = True) -> None:
    parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{radius}" fill="{fill}" stroke="{stroke}" stroke-width="2"/>')
    gap = font_size * 1.35
    start = y + h / 2 - gap * (len(lines) - 1) / 2 + font_size * 0.34
    for i, line in enumerate(lines):
        weight = "600" if bold_first and i == 0 else "400"
        parts.append(
            f'<text x="{x + w / 2}" y="{start + i * gap}" text-anchor="middle" '
            f'font-family="{FONT_CN}" font-size="{font_size}" font-weight="{weight}" fill="#111111">{esc(line)}</text>'
        )


def representation_box(parts: list[str], x: float, y: float, w: float, h: float, title: str, subscript: str, fill: str) -> None:
    parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="8" fill="{fill}" stroke="{STROKE}" stroke-width="2"/>')
    parts.append(
        f'<text x="{x + w / 2}" y="{y + h / 2 - 10}" text-anchor="middle" '
        f'font-family="{FONT_CN}" font-size="26" font-weight="600" fill="#111111">{esc(title)}</text>'
    )
    parts.append(
        f'<text x="{x + w / 2}" y="{y + h / 2 + 31}" text-anchor="middle" '
        f'font-family="{FONT_EN}" font-size="26" font-style="italic" fill="#111111">'
        f'z<tspan baseline-shift="sub" font-size="18">{esc(subscript)}</tspan></text>'
    )


def symbol_box(parts: list[str], x: float, y: float, w: float, h: float, title: str, base: str, subscript: str, suffix: str, fill: str) -> None:
    parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="8" fill="{fill}" stroke="{STROKE}" stroke-width="2"/>')
    parts.append(
        f'<text x="{x + w / 2}" y="{y + h / 2 - 10}" text-anchor="middle" '
        f'font-family="{FONT_CN}" font-size="25" font-weight="600" fill="#111111">{esc(title)}</text>'
    )
    parts.append(
        f'<text x="{x + w / 2}" y="{y + h / 2 + 30}" text-anchor="middle" '
        f'font-family="{FONT_EN}" font-size="25" font-style="italic" fill="#111111">'
        f'{esc(base)}<tspan baseline-shift="sub" font-size="17">{esc(subscript)}</tspan>'
        f'<tspan baseline-shift="baseline" font-size="25">{esc(suffix)}</tspan></text>'
    )


def label(parts: list[str], x: float, y: float, text: str, *, size: int = 25, weight: int = 400, anchor: str = "middle", family: str = FONT_CN) -> None:
    parts.append(f'<text x="{x}" y="{y}" text-anchor="{anchor}" font-family="{family}" font-size="{size}" font-weight="{weight}" fill="#111111">{esc(text)}</text>')


def arrow(parts: list[str], x1: float, y1: float, x2: float, y2: float) -> None:
    parts.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{STROKE}" stroke-width="2.2" marker-end="url(#arrow)"/>')


def poly_arrow(parts: list[str], points: list[tuple[float, float]]) -> None:
    p = ' '.join(f'{x},{y}' for x, y in points)
    parts.append(f'<polyline points="{p}" fill="none" stroke="{STROKE}" stroke-width="2.2" marker-end="url(#arrow)"/>')


def divider(parts: list[str], x1: float, y1: float, x2: float, y2: float) -> None:
    parts.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="#777777" stroke-width="1.6"/>')


def figure_1_1() -> str:
    h = 760
    p = svg_start(h)
    box(p, 630, 25, 340, 72, ["第1章  引言"], fill=GRAY)
    arrow(p, 800, 97, 800, 132)
    box(p, 630, 140, 340, 72, ["第2章  相关理论与关键技术"], fill=GRAY, font_size=25)
    arrow(p, 800, 212, 800, 265)

    p.append(
        '<rect x="55" y="275" width="1490" height="245" rx="16" '
        'fill="#FFFFFF" stroke="#777777" stroke-width="1.8" stroke-dasharray="10 8"/>'
    )
    p.append('<rect x="675" y="258" width="250" height="38" fill="#FFFFFF"/>')
    label(p, 800, 285, "核心方法研究", size=26, weight=600)

    box(p, 115, 335, 390, 125, ["第3章", "三模态可信度评估"], fill=API, font_size=26)
    box(p, 605, 335, 390, 125, ["第4章", "可信证据融合"], fill=GRAPH, font_size=26)
    box(p, 1095, 335, 390, 125, ["第5章", "恶意漏报风险控制"], fill=MAN, font_size=26)
    arrow(p, 505, 397.5, 595, 397.5)
    arrow(p, 995, 397.5, 1085, 397.5)

    arrow(p, 800, 520, 800, 555)
    box(p, 630, 565, 340, 72, ["第6章  实验设计与结果分析"], fill=GRAY, font_size=25)
    arrow(p, 800, 637, 800, 670)
    box(p, 630, 678, 340, 62, ["第7章  总结与展望"], fill=GRAY, font_size=25)
    return svg_end(p)


def figure_2_1() -> str:
    h = 680
    p = svg_start(h)
    box(p, 35, 270, 240, 120, ["APK文件", "代码与配置"], fill=GRAY, font_size=26)
    box(p, 370, 105, 280, 120, ["DEX字节码", "方法与调用"], fill=GRAY, font_size=26)
    box(p, 370, 502, 280, 120, ["Manifest文件", "权限与组件"], fill=GRAY, font_size=26)

    box(p, 800, 55, 300, 125, ["API调用序列", "行为顺序信息"], fill=API, font_size=26)
    box(p, 800, 278, 300, 125, ["程序调用图", "程序结构信息"], fill=GRAPH, font_size=26)
    box(p, 800, 500, 300, 125, ["Manifest特征", "应用声明信息"], fill=MAN, font_size=26)

    box(p, 1300, 270, 250, 145, ["三模态输入", "行为、结构与声明"], fill=GRAY, font_size=26)

    poly_arrow(p, [(275, 330), (325, 330), (325, 165), (360, 165)])
    poly_arrow(p, [(275, 330), (325, 330), (325, 562), (360, 562)])
    poly_arrow(p, [(650, 165), (720, 165), (720, 117), (790, 117)])
    poly_arrow(p, [(650, 165), (720, 165), (720, 340), (790, 340)])
    arrow(p, 650, 562, 790, 562)

    arrow(p, 1100, 117, 1190, 117)
    arrow(p, 1100, 340, 1190, 340)
    arrow(p, 1100, 562, 1190, 562)
    divider(p, 1200, 117, 1200, 562)
    arrow(p, 1200, 340, 1290, 340)
    return svg_end(p)


def figure_2_2() -> str:
    h = 720
    p = svg_start(h)
    centers = [190, 580, 970, 1360]
    for x, text in zip(centers, ["输入模态", "初始表示", "编码器", "模态表示"]):
        label(p, x, 65, text, weight=600)
    rows = [
        (145, "API调用序列", API, "词向量与位置编码", "Transformer编码", "行为序列表示", "api"),
        (355, "程序调用图", GRAPH, "节点特征映射", "GATv2编码", "程序结构表示", "graph"),
        (565, "Manifest特征", MAN, "特征映射", "多层感知机编码", "应用声明表示", "man"),
    ]
    xs = [30, 420, 810, 1200]
    bw = 320
    for cy, src, color, initial, encoder, output, subscript in rows:
        box(p, xs[0], cy - 65, bw, 130, [src], fill=color, font_size=26)
        box(p, xs[1], cy - 65, bw, 130, [initial], fill="#FFFFFF", font_size=26)
        box(p, xs[2], cy - 65, bw, 130, [encoder], fill="#FFFFFF", font_size=26)
        representation_box(p, xs[3], cy - 65, bw, 130, output, subscript, color)
        arrow(p, xs[0] + bw, cy, xs[1] - 10, cy)
        arrow(p, xs[1] + bw, cy, xs[2] - 10, cy)
        arrow(p, xs[2] + bw, cy, xs[3] - 10, cy)
    return svg_end(p)


def figure_2_3() -> str:
    h = 500
    p = svg_start(h)
    steps = [
        (25, 160, 270, 170, ["三模态表示", "行为、结构与声明"], GRAY),
        (335, 160, 270, 170, ["模态可信度评估", "样本级可信评分"], GRAY),
        (645, 160, 270, 170, ["证据表示与折扣", "类别支持与不确定性"], GRAY),
        (955, 160, 270, 170, ["冲突保留融合", "类别概率与不确定性"], GRAY),
        (1265, 160, 270, 170, ["共形预测与拒识", "接受或转交复核"], GRAY),
    ]
    for i, (x, y, w, hh, lines, fill) in enumerate(steps):
        box(p, x, y, w, hh, lines, fill=fill, font_size=26)
        if i < len(steps) - 1:
            nx = steps[i + 1][0]
            arrow(p, x + w, y + hh / 2, nx - 10, y + hh / 2)
    return svg_end(p)


def figure_3_1() -> str:
    h = 650
    p = svg_start(h)
    label(p, 150, 55, "静态模态", weight=600)
    label(p, 500, 55, "可观测证据", weight=600)
    label(p, 830, 55, "可靠性估计", weight=600)
    label(p, 1160, 55, "综合修正", weight=600)
    label(p, 1450, 55, "可信评分", weight=600)

    modality_rows = [
        (145, "API调用序列", API),
        (325, "程序调用图", GRAPH),
        (505, "Manifest声明", MAN),
    ]
    for cy, text, color in modality_rows:
        box(p, 30, cy - 55, 240, 110, [text], fill=color, font_size=25)
        arrow(p, 270, cy, 320, cy)

    divider(p, 330, 145, 330, 505)
    arrow(p, 330, 325, 350, 325)

    box(p, 360, 220, 280, 210, ["按模态构建证据", "模态自身质量", "跨模态支持与一致性"], fill="#FFFFFF", font_size=24)
    box(p, 710, 245, 250, 160, ["单调可靠性校准", "分支正确性监督"], fill=GRAY, font_size=24)
    box(p, 1030, 245, 260, 160, ["综合可信度因素", "模态可用性", "能力先验与可见修正"], fill=GRAY, font_size=23)
    box(p, 1360, 260, 210, 130, ["模态可信评分", "Tₘ(x)"], fill=ACCENT, font_size=25)
    arrow(p, 640, 325, 700, 325)
    arrow(p, 960, 325, 1020, 325)
    arrow(p, 1290, 325, 1350, 325)
    return svg_end(p)


def figure_3_2() -> str:
    h = 700
    p = svg_start(h)
    headers = [(165, "输入信息"), (535, "计算过程"), (880, "中间结果"), (1190, "组合"), (1460, "输出")]
    for x, text in headers:
        label(p, x, 52, text, weight=600)

    rows = [
        (125, ["可观测证据", "oₘ(x)"], ["单调可靠性校准"], ["样本级可靠性", "Rₘ(x)"]),
        (280, ["校准集分支性能"], ["相对性能归一化"], ["能力先验", "Cₘ"]),
        (435, ["完整性与可见比例", "正常可见水平"], ["相对可见性修正"], ["可见修正", "M̃ₘ(x)"]),
        (590, ["解析状态与有效内容"], ["模态可用性判断"], ["模态可用性", "aₘ(x)"]),
    ]
    for cy, inp, process, result in rows:
        box(p, 30, cy - 55, 280, 110, inp, fill="#FFFFFF", font_size=23)
        box(p, 400, cy - 55, 270, 110, process, fill=GRAY, font_size=23)
        box(p, 760, cy - 55, 240, 110, result, fill="#FFFFFF", font_size=23)
        arrow(p, 310, cy, 390, cy)
        arrow(p, 670, cy, 750, cy)
        arrow(p, 1000, cy, 1060, cy)

    divider(p, 1070, 125, 1070, 590)
    arrow(p, 1070, 357, 1100, 357)
    box(p, 1110, 295, 180, 125, ["乘性组合"], fill=GRAY, font_size=25)
    arrow(p, 1290, 357, 1330, 357)
    box(p, 1340, 285, 230, 145, ["模态可信评分", "Tₘ(x)"], fill=ACCENT, font_size=25)
    return svg_end(p)


def figure_4_1() -> str:
    h = 700
    p = svg_start(h)
    headers = [(140, "原始证据意见"), (415, "可信度折扣"), (700, "折扣后意见"), (1030, "联合计算"), (1420, "结果输出")]
    for x, text in headers:
        label(p, x, 50, text, weight=600)

    rows = [
        (125, "API证据意见", API),
        (330, "调用图证据意见", GRAPH),
        (535, "Manifest证据意见", MAN),
    ]
    for cy, title, color in rows:
        symbol_box(p, 25, cy - 60, 230, 120, title, "ω", "m", "", color)
        symbol_box(p, 300, cy - 60, 230, 120, "可信度折扣", "T", "m", "(x)", "#FFFFFF")
        symbol_box(p, 575, cy - 60, 250, 120, "折扣后意见", "ω̃", "m", "", color)
        arrow(p, 255, cy, 290, cy)
        arrow(p, 530, cy, 565, cy)
        arrow(p, 825, cy, 870, cy)

    divider(p, 880, 125, 880, 535)
    arrow(p, 880, 330, 910, 330)
    box(p, 920, 245, 270, 170, ["三模态联合计算", "类别共同支持", "模态冲突计算"], fill=GRAY, font_size=24)

    poly_arrow(p, [(1190, 300), (1230, 300), (1230, 197), (1270, 197)])
    box(p, 1280, 135, 270, 125, ["组合前模态冲突", "κ"], fill="#FFFFFF", font_size=24)

    arrow(p, 1055, 415, 1055, 470)
    box(p, 920, 480, 270, 145, ["条件意见路由", "π分配与argmax风险"], fill=GRAY, font_size=24)
    arrow(p, 1190, 552, 1270, 552)
    box(p, 1280, 480, 270, 145, ["融合输出", "pᶠ，uᶠ"], fill=ACCENT, font_size=24)
    return svg_end(p)


def figure_4_2_optional() -> str:
    h = 600
    p = svg_start(h)
    box(p, 70, 215, 260, 150, ["多模态证据", "共同支持与冲突"], fill=GRAY)
    arrow(p, 330, 290, 455, 290)
    box(p, 470, 115, 330, 150, ["Dempster规则", "归一化非冲突证据"], fill="#FFFFFF")
    box(p, 470, 355, 330, 150, ["条件意见路由", "学习样本级路由权重"], fill="#FFFFFF")
    poly_arrow(p, [(400, 290), (400, 190), (460, 190)])
    poly_arrow(p, [(400, 290), (400, 430), (460, 430)])
    box(p, 920, 115, 330, 150, ["类别支持增强", "冲突影响被弱化"], fill=GRAY)
    box(p, 920, 355, 330, 150, ["不确定性提高", "保留模态分歧"], fill=GRAY)
    arrow(p, 800, 190, 910, 190)
    arrow(p, 800, 430, 910, 430)
    box(p, 1360, 355, 180, 150, ["本文采用", "条件意见路由"], fill=ACCENT)
    arrow(p, 1250, 430, 1350, 430)
    return svg_end(p)


def figure_5_1() -> str:
    h = 700
    p = svg_start(h)
    label(p, 35, 58, "校准阶段", size=27, weight=600, anchor="start")
    calibration = [
        (90, 80, 260, ["独立校准集"]),
        (430, 80, 270, ["类别非一致性分数", "r_c(x)"]),
        (790, 80, 260, ["类别条件阈值", "τ₀，τ₁"]),
    ]
    for i, (x, y, w, lines) in enumerate(calibration):
        box(p, x, y, w, 110, lines, fill=GRAY, font_size=25)
        if i < len(calibration) - 1:
            arrow(p, x + w, 135, calibration[i + 1][0] - 10, 135)

    label(p, 35, 302, "测试阶段", size=27, weight=600, anchor="start")
    testing = [
        (90, 325, 260, ["融合输出", "pᶠ(x)，κ(x)"]),
        (430, 325, 270, ["测试样本类别分数", "r₀(x)，r₁(x)"]),
        (790, 325, 260, ["类别阈值检验", "τ₀，τ₁"]),
        (1130, 325, 270, ["共形预测集合", "Γ(x)"]),
    ]
    for i, (x, y, w, lines) in enumerate(testing):
        box(p, x, y, w, 120, lines, fill="#FFFFFF" if i < 3 else GRAY, font_size=25)
        if i < len(testing) - 1:
            arrow(p, x + w, 385, testing[i + 1][0] - 10, 385)

    arrow(p, 920, 190, 920, 315)

    outcomes = [
        (735, ["|Γ(x)|=1", "接受预测"], API),
        (1030, ["|Γ(x)|=0", "拒识：证据不足"], REJECT),
        (1325, ["|Γ(x)|=2", "拒识：类别歧义"], REJECT),
    ]
    for x, lines, fill in outcomes:
        box(p, x, 545, 240, 110, lines, fill=fill, font_size=23)

    divider(p, 1265, 445, 1265, 500)
    poly_arrow(p, [(1265, 500), (855, 500), (855, 535)])
    poly_arrow(p, [(1265, 500), (1150, 500), (1150, 535)])
    poly_arrow(p, [(1265, 500), (1445, 500), (1445, 535)])
    return svg_end(p)


FIGURES = {
    "图1-1_论文组织结构": figure_1_1,
    "图2-1_模态特征提取示意图": figure_2_1,
    "图2-2_三模态特征编码过程": figure_2_2,
    "图2-3_多模态可信融合基本过程": figure_2_3,
    "图3-1_三模态可观测证据与可信度评估框架": figure_3_1,
    "图3-2_三模态可信度评估流程": figure_3_2,
    "图4-1_三模态证据调整与冲突保留融合框架": figure_4_1,
    "图4-2_冲突处理方式对比_可选": figure_4_2_optional,
    "图5-1_类别条件共形预测与拒识流程": figure_5_1,
}


CAPTIONS = {
    "图1-1_论文组织结构": ("图1-1 论文组织结构", "Figure 1-1 Organization of the Thesis"),
    "图2-1_模态特征提取示意图": ("图2-1 模态特征提取示意图", "Figure 2-1 Static Extraction of the Three Modalities"),
    "图2-2_三模态特征编码过程": ("图2-2 三模态特征编码过程", "Figure 2-2 Encoding Process of the Three Modalities"),
    "图2-3_多模态可信融合基本过程": ("图2-3 多模态可信融合基本过程", "Figure 2-3 Basic Process of Trustworthy Multimodal Fusion"),
    "图3-1_三模态可观测证据与可信度评估框架": ("图3-1 三模态可观测证据与可信度评估框架", "Figure 3-1 Observable Evidence and Trustworthiness Assessment Framework"),
    "图3-2_三模态可信度评估流程": ("图3-2 三模态可信度评估流程", "Figure 3-2 Trustworthiness Assessment Process of the Three Modalities"),
    "图4-1_三模态证据调整与冲突保留融合框架": ("图4-1 三模态证据调整与冲突保留融合框架", "Figure 4-1 Evidence Adjustment and Conflict-preserving Fusion Framework"),
    "图4-2_冲突处理方式对比_可选": ("图4-2 Dempster融合与条件意见路由对比", "Figure 4-2 Comparison of Dempster Fusion and Conditional Opinion Routing"),
    "图5-1_类别条件共形预测与拒识流程": ("图5-1 类别条件共形预测与拒识流程", "Figure 5-1 Class-conditional Conformal Prediction and Rejection Process"),
}


def write_svgs() -> None:
    SVG_DIR.mkdir(parents=True, exist_ok=True)
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    for name, fn in FIGURES.items():
        (SVG_DIR / f"{name}.svg").write_text(fn(), encoding="utf-8")


def svg_to_png() -> None:
    if not EDGE.exists():
        raise FileNotFoundError(EDGE)
    for svg_path in SVG_DIR.glob("*.svg"):
        out = PNG_DIR / f"{svg_path.stem}.png"
        svg_text = svg_path.read_text(encoding="utf-8")
        height_match = re.search(r'<svg[^>]+height="(\d+)"', svg_text)
        target_height = int(height_match.group(1)) if height_match else 900
        uri = svg_path.resolve().as_uri()
        subprocess.run(
            [
                str(EDGE),
                "--headless",
                "--disable-gpu",
                "--hide-scrollbars",
                "--force-device-scale-factor=1",
                "--window-size=1600,900",
                f"--screenshot={out}",
                uri,
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        with Image.open(out) as image:
            image.crop((0, 0, W, target_height)).save(out)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge, val, sz in (("top", "single", "12"), ("bottom", "single", "12"), ("insideH", "nil", "0"), ("left", "nil", "0"), ("right", "nil", "0"), ("insideV", "nil", "0")):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), "333333")


def set_cell_bottom_border(cell, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:color"), "333333")


def style_table(table, widths_cm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            set_repeat_table_header(row)
        for col_idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, "F2F2F2")
                set_cell_bottom_border(cell)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 or row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.15
                for run in para.runs:
                    run.font.name = "SimSun"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9.5)
                    if row_idx == 0:
                        run.bold = True


def caption(doc: Document, zh: str, en: str, before: bool = False) -> None:
    for text, font, size in ((zh, "SimSun", 10.5), (en, "Times New Roman", 10.0)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3 if before else 0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = font
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体" if font == "SimSun" else "Times New Roman")
        r.font.size = Pt(size)


def add_table_3_1(doc: Document) -> None:
    caption(doc, "表3-1 三模态可信度评估中的可观测证据", "Table 3-1 Observable Evidence for Three-modality Trustworthiness Assessment", before=True)
    rows = [
        ["证据", "符号", "主要来源", "作用"],
        ["模态可用性", "aₘ(x)", "解析状态与有效内容", "排除不可用模态"],
        ["模态完整性", "qₘ(x)", "解析、保留与特征有效性", "描述模态信息质量"],
        ["模型可见比例", "vₘ(x)", "编码器实际接收范围", "刻画输入预算影响"],
        ["模型可见完整性", "eₘ(x)", "完整性与可见比例", "衡量有效可见信息"],
        ["结构支持程度", "sₐg(x)", "API事件与方法节点对应", "描述代码侧结构支持"],
        ["语义支持程度", "sₘc(x)", "声明与代码安全语义", "描述跨源语义支持"],
        ["双向不一致程度", "dₘc(x), dcm(x)", "声明与代码语义差异", "描述跨源不一致"],
        ["分支能力先验", "Cₘ", "校准集分支性能", "修正分支整体能力差异"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    style_table(table, [3.2, 2.4, 5.0, 4.8])


def add_blank_experiment_tables(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph("第6章三线表模板（数值留空，填写后使用）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    specs = [
        ("表6-1 数据集划分与类别分布", "Table 6-1 Dataset Splits and Class Distribution", ["数据划分", "良性样本", "恶意样本", "总数"], [3.8, 3.8, 3.8, 3.8], 4),
        ("表6-2 主要训练参数", "Table 6-2 Main Training Parameters", ["参数", "取值", "参数", "取值"], [4.5, 3.0, 4.5, 3.0], 7),
        ("表6-3 不同方法的常规检测结果", "Table 6-3 Detection Results of Different Methods", ["方法", "Macro-F1", "AUC", "AP", "Brier", "ECE"], [4.2, 2.2, 2.0, 2.0, 2.2, 2.2], 8),
        ("表6-4 不同可信融合方法的对比结果", "Table 6-4 Comparison of Trustworthy Fusion Methods", ["方法", "Macro-F1", "AURC", "Brier", "ECE"], [5.4, 2.4, 2.4, 2.4, 2.4], 7),
        ("表6-5 关键模块分析结果", "Table 6-5 Analysis Results of Key Modules", ["实验设置", "Macro-F1", "AURC", "ECE", "说明"], [4.7, 2.2, 2.2, 2.2, 4.3], 7),
        ("表6-6 自然低可信子集检测结果", "Table 6-6 Detection Results on Naturally Low-confidence Subsets", ["方法", "低完整性", "低支持", "高冲突", "低接受度"], [4.4, 2.7, 2.7, 2.7, 2.7], 7),
        ("表6-7 恶意漏报风险控制结果", "Table 6-7 Results of Malware False-Negative Risk Control", ["方法", "接受率", "经验覆盖率", "选择性风险", "恶意漏报数"], [4.4, 2.6, 3.0, 3.0, 3.0], 6),
    ]
    for spec_idx, (zh, en, headers, widths, row_count) in enumerate(specs):
        if spec_idx > 0:
            doc.add_page_break()
        caption(doc, zh, en, before=True)
        table = doc.add_table(rows=row_count, cols=len(headers))
        for j, text in enumerate(headers):
            table.cell(0, j).text = text
        for i in range(1, row_count):
            for j in range(len(headers)):
                table.cell(i, j).text = "—"
        style_table(table, widths)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)


def build_docx() -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.5)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("定稿13统一图表替换稿")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "SimSun"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    note = doc.add_paragraph("说明：图形采用统一配色、线宽和字号。复制到论文时优先插入SVG；PNG用于兼容。图题保留为Word文本，不嵌入图片。")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(12)
    for name in FIGURES:
        if "可选" in name:
            continue
        png = PNG_DIR / f"{name}.png"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(png), width=Cm(15.0))
        zh, en = CAPTIONS[name]
        caption(doc, zh, en)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    doc.add_page_break()
    optional = "图4-2_冲突处理方式对比_可选"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(PNG_DIR / f"{optional}.png"), width=Cm(15.0))
    caption(doc, *CAPTIONS[optional])
    add_table_3_1(doc)
    add_blank_experiment_tables(doc)
    out = ROOT / "定稿13_统一图表替换稿.docx"
    doc.save(out)
    return out


def write_readme() -> None:
    lines = [
        "# 定稿13统一图表替换包",
        "",
        "- `svg/`：可编辑矢量图，论文中优先使用。",
        "- `png/`：兼容版高分辨率图片。",
        "- `定稿13_统一图表替换稿.docx`：全部图、表3-1和第6章三线表模板。",
        "- 图题应保留为Word段落，不要写入图片。",
        "- 图4-2为可选补充图，未自动改变正文编号。",
    ]
    (ROOT / "使用说明.md").write_text('\n'.join(lines), encoding="utf-8")


def main() -> None:
    write_svgs()
    svg_to_png()
    build_docx()
    write_readme()
    print(ROOT)


if __name__ == "__main__":
    main()
